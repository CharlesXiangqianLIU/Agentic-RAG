# ingestion/parser.py
"""
Parsing surface for the knowledge-rag ingestion pipeline.

``parse_document(path)`` is the canonical entry point — it dispatches by
file extension to the matching backend (``parse_docx``, ``parse_markdown``,
``parse_text``, ``parse_pdf``, ``parse_html``). All backends return a
``ParsedDocument`` with the same shape, so downstream chunking / indexing
code is agnostic to the source format.

``parse_docx`` (defined here) is the .docx-specific implementation; the
remaining backends live in ``ingestion/parsers/``.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class ParsedParagraph:
    text: str
    page_number: int
    is_header: bool


@dataclass
class ParsedRow:
    cells: list[str]
    page_number: int

    def to_text(self) -> str:
        return " | ".join(c.strip() for c in self.cells)


@dataclass
class ParsedTable:
    rows: list[ParsedRow]
    section: str          # section heading active when the table was encountered
    page_number: int


@dataclass
class ParsedDocument:
    filename: str
    paragraphs: list[ParsedParagraph] = field(default_factory=list)
    tables: list[ParsedTable] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_HEADING_STYLE_RE = re.compile(r"^[Hh]eading\s*[1-6]$")


def _is_heading_style(style_name: str | None) -> bool:
    if style_name is None:
        return False
    return bool(_HEADING_STYLE_RE.match(style_name))


def _looks_like_header(text: str) -> bool:
    """Heuristic: short ALL-CAPS or title-case line with no period."""
    t = text.strip()
    if not t or len(t) > 120:
        return False
    if t.endswith("."):
        return False
    # ALL-CAPS words (ignoring numbers/punctuation)
    words = re.findall(r"[A-Za-z]+", t)
    if words and all(w.isupper() for w in words) and len(words) <= 12:
        return True
    return False


def _count_page_breaks_in_para(para_elem) -> int:
    """Count explicit and rendered page-breaks inside a <w:p> element."""
    count = 0
    for br in para_elem.iter(qn("w:br")):
        typ = br.get(qn("w:type"), "")
        if typ == "page":
            count += 1
    for _ in para_elem.iter(qn("w:lastRenderedPageBreak")):
        count += 1
    return count


def _para_text(para_elem) -> str:
    """Extract plain text from a <w:p> element."""
    parts = []
    for t in para_elem.iter(qn("w:t")):
        parts.append(t.text or "")
    return "".join(parts)


def _row_cells_text(row_elem) -> list[str]:
    """Return list of cell texts for a <w:tr> element."""
    cells = []
    for tc in row_elem.findall(".//" + qn("w:tc")):
        texts = []
        for t in tc.iter(qn("w:t")):
            texts.append(t.text or "")
        cells.append("".join(texts))
    return cells


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def parse_docx(path: Path | str) -> ParsedDocument:
    """Parse a .docx file and return a ParsedDocument."""
    path = Path(path)
    doc = Document(str(path))
    result = ParsedDocument(filename=path.name)

    page = 1
    current_section = "Introduction"
    body = doc.element.body

    for elem in body:
        tag = elem.tag

        # ── Paragraph ────────────────────────────────────────────────────
        if tag == qn("w:p"):
            page += _count_page_breaks_in_para(elem)

            text = _para_text(elem).strip()
            if not text:
                continue

            # Determine heading status
            style_name: str | None = None
            pPr = elem.find(qn("w:pPr"))
            if pPr is not None:
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is not None:
                    style_name = pStyle.get(qn("w:val"))

            is_hdr = _is_heading_style(style_name) or _looks_like_header(text)
            if is_hdr:
                current_section = text

            result.paragraphs.append(ParsedParagraph(
                text=text,
                page_number=page,
                is_header=is_hdr,
            ))

        # ── Table ─────────────────────────────────────────────────────────
        elif tag == qn("w:tbl"):
            rows: list[ParsedRow] = []
            for tr in elem.findall(".//" + qn("w:tr")):
                cells = _row_cells_text(tr)
                if any(c.strip() for c in cells):       # skip blank rows
                    rows.append(ParsedRow(cells=cells, page_number=page))

            if rows:
                result.tables.append(ParsedTable(
                    rows=rows,
                    section=current_section,
                    page_number=page,
                ))

    return result


# ---------------------------------------------------------------------------
# Format-agnostic entry point
# ---------------------------------------------------------------------------


def parse_document(path: Path | str) -> ParsedDocument:
    """Parse a document of any supported format into a ``ParsedDocument``.

    Supported extensions: ``.docx``, ``.md``, ``.markdown``, ``.txt``,
    ``.pdf``, ``.html``, ``.htm``. Other extensions raise ``ValueError``.
    """
    p = Path(path)
    ext = p.suffix.lower()
    if ext == ".docx":
        return parse_docx(p)
    if ext in (".md", ".markdown"):
        from ingestion.parsers.markdown import parse_markdown
        return parse_markdown(p)
    if ext == ".txt":
        from ingestion.parsers.text import parse_text
        return parse_text(p)
    if ext == ".pdf":
        from ingestion.parsers.pdf import parse_pdf
        return parse_pdf(p)
    if ext in (".html", ".htm"):
        from ingestion.parsers.html import parse_html
        return parse_html(p)
    raise ValueError(f"Unsupported document extension: {ext!r} ({p.name})")


SUPPORTED_EXTENSIONS = (".docx", ".md", ".markdown", ".txt", ".pdf", ".html", ".htm")
