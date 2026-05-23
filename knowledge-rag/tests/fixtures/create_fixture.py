#!/usr/bin/env python3
"""Generate a domain-neutral sample.docx test fixture.

The document mirrors the structure expected by ``test_parser.py`` and
``test_chunker.py``: an H1 heading, a normal paragraph, a 6-column table
with one header row and two data rows, an H2 heading, and a closing
paragraph. The content is generic quarterly-business data so the tests
do not depend on any domain pack being active.
"""

from pathlib import Path
from docx import Document

doc = Document()

doc.add_heading("Quarterly Performance Review", level=1)
doc.add_paragraph(
    "The team reviewed all open accounts and recorded the headline numbers below."
)

table = doc.add_table(rows=3, cols=6)
table.style = "Light Grid Accent 1"

headers = ["Account", "Region", "Quarter", "Status", "Owner", "Revenue"]
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header

row1 = ["A-001", "North", "Q1 2026", "Active", "Alice", "120000"]
for i, value in enumerate(row1):
    table.rows[1].cells[i].text = value

row2 = ["A-002", "South", "Q1 2026", "Active", "Bob", "150000"]
for i, value in enumerate(row2):
    table.rows[2].cells[i].text = value

doc.add_heading("Notes", level=2)
doc.add_paragraph(
    "Account A-002 grew the fastest this quarter; A-001 stayed on plan."
)

output_path = Path(__file__).parent / "sample.docx"
doc.save(str(output_path))
print(f"Successfully created {output_path}")
